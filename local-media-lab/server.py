from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import threading
import time
import uuid
import webbrowser
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
import uvicorn

APP_NAME = "Open Social Scheduler - Media Lab Locale"
APP_VERSION = "1.0.0"
ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"
INCOMING_DIR = ROOT / "incoming"
OUTPUT_DIR = ROOT / "elaborazioni"
SETTINGS_FILE = ROOT / "settings.json"
PID_FILE = ROOT / "media_lab.pid"
HOST = "127.0.0.1"
PORT = 8765

ALLOWED_EXTENSIONS = {
    ".mp4", ".mov", ".mkv", ".avi", ".webm", ".m4v", ".mpeg", ".mpg",
    ".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus",
}
AUDIO_FORMATS = {"mp3", "wav", "m4a"}
WHISPER_MODELS = {"tiny", "base", "small", "medium"}
LANGUAGES = {"it", "auto"}

app = FastAPI(title=APP_NAME, version=APP_VERSION)
app.mount("/assets", StaticFiles(directory=STATIC_DIR), name="assets")

_jobs: dict[str, dict[str, Any]] = {}
_jobs_lock = threading.Lock()
_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="media-lab")
_whisper_models: dict[str, Any] = {}
_whisper_lock = threading.Lock()


def ensure_directories() -> None:
    STATIC_DIR.mkdir(parents=True, exist_ok=True)
    INCOMING_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def read_settings() -> dict[str, Any]:
    defaults = {
        "output_dir": str(OUTPUT_DIR),
        "keep_original": False,
        "default_audio_format": "mp3",
        "default_whisper_model": "base",
        "default_language": "it",
    }
    if not SETTINGS_FILE.exists():
        return defaults
    try:
        data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return defaults
        return {**defaults, **data}
    except Exception:
        return defaults


def write_settings(settings: dict[str, Any]) -> None:
    SETTINGS_FILE.write_text(
        json.dumps(settings, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def output_root() -> Path:
    configured = Path(str(read_settings().get("output_dir", OUTPUT_DIR))).expanduser()
    configured.mkdir(parents=True, exist_ok=True)
    return configured.resolve()


def safe_filename(name: str) -> str:
    base = Path(name or "file").name
    stem = Path(base).stem
    suffix = Path(base).suffix.lower()
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._") or "file"
    return f"{clean[:120]}{suffix}"


def safe_stem(name: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("._")
    return clean[:120] or "media"


def ffmpeg_path() -> str | None:
    candidates = [
        shutil.which("ffmpeg"),
        str(ROOT / "ffmpeg" / "bin" / "ffmpeg.exe"),
        str(ROOT / "bin" / "ffmpeg.exe"),
        str(ROOT / "ffmpeg.exe"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def whisper_available() -> bool:
    try:
        import whisper  # noqa: F401
        return True
    except Exception:
        return False


def update_job(job_id: str, **changes: Any) -> None:
    with _jobs_lock:
        if job_id in _jobs:
            _jobs[job_id].update(changes)
            _jobs[job_id]["updated_at"] = time.time()


def get_job(job_id: str) -> dict[str, Any]:
    with _jobs_lock:
        job = _jobs.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Elaborazione non trovata")
        return dict(job)


def add_log(job_id: str, message: str) -> None:
    with _jobs_lock:
        if job_id not in _jobs:
            return
        logs = list(_jobs[job_id].get("logs", []))
        logs.append(message)
        _jobs[job_id]["logs"] = logs[-30:]
        _jobs[job_id]["updated_at"] = time.time()


def run_command(command: list[str], job_id: str) -> None:
    add_log(job_id, "Esecuzione FFmpeg in corso...")
    creationflags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    process = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        errors="replace",
        creationflags=creationflags,
    )
    if process.returncode != 0:
        error = process.stderr.strip()[-3000:] or "Errore FFmpeg senza dettagli"
        raise RuntimeError(error)


def audio_command(source: Path, destination: Path, audio_format: str) -> list[str]:
    ffmpeg = ffmpeg_path()
    if not ffmpeg:
        raise RuntimeError(
            "FFmpeg non è installato o non è raggiungibile. Esegui INSTALLA_MEDIA_LAB.bat."
        )
    base = [ffmpeg, "-hide_banner", "-loglevel", "error", "-y", "-i", str(source), "-vn"]
    if audio_format == "mp3":
        return base + ["-codec:a", "libmp3lame", "-b:a", "192k", str(destination)]
    if audio_format == "wav":
        return base + ["-codec:a", "pcm_s16le", "-ar", "44100", "-ac", "2", str(destination)]
    if audio_format == "m4a":
        return base + ["-codec:a", "aac", "-b:a", "192k", str(destination)]
    raise RuntimeError("Formato audio non supportato")


def whisper_input_command(source: Path, destination: Path) -> list[str]:
    ffmpeg = ffmpeg_path()
    if not ffmpeg:
        raise RuntimeError("FFmpeg non disponibile")
    return [
        ffmpeg, "-hide_banner", "-loglevel", "error", "-y", "-i", str(source), "-vn",
        "-codec:a", "pcm_s16le", "-ar", "16000", "-ac", "1", str(destination),
    ]


def load_whisper_model(model_name: str) -> Any:
    with _whisper_lock:
        if model_name in _whisper_models:
            return _whisper_models[model_name]
        import whisper
        model = whisper.load_model(model_name)
        _whisper_models[model_name] = model
        return model


def srt_timestamp(seconds: float) -> str:
    milliseconds = max(0, int(round(seconds * 1000)))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def write_transcription_files(result: dict[str, Any], output_dir: Path, base_name: str) -> list[Path]:
    text = str(result.get("text", "")).strip()
    txt_path = output_dir / f"{base_name}_trascrizione.txt"
    txt_path.write_text(text + ("\n" if text else ""), encoding="utf-8")

    srt_path = output_dir / f"{base_name}_sottotitoli.srt"
    segments = result.get("segments") or []
    srt_lines: list[str] = []
    for index, segment in enumerate(segments, start=1):
        start = float(segment.get("start", 0.0))
        end = float(segment.get("end", start))
        segment_text = str(segment.get("text", "")).strip()
        srt_lines.extend([
            str(index),
            f"{srt_timestamp(start)} --> {srt_timestamp(end)}",
            segment_text,
            "",
        ])
    srt_path.write_text("\n".join(srt_lines), encoding="utf-8")

    docx_path: Path | None = output_dir / f"{base_name}_trascrizione.docx"
    try:
        from docx import Document
        document = Document()
        document.add_heading("Trascrizione audio", level=1)
        if text:
            for paragraph in [part.strip() for part in text.split("\n") if part.strip()]:
                document.add_paragraph(paragraph)
        else:
            document.add_paragraph("Nessun testo riconosciuto.")
        document.save(docx_path)
    except Exception:
        docx_path = None

    paths = [txt_path, srt_path]
    if docx_path is not None and docx_path.exists():
        paths.append(docx_path)
    return paths


def artifact_label(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".mp3", ".wav", ".m4a"}:
        return "Audio estratto"
    if suffix == ".txt":
        return "Trascrizione TXT"
    if suffix == ".docx":
        return "Trascrizione Word"
    if suffix == ".srt":
        return "Sottotitoli SRT"
    return path.name


def serialize_artifacts(job_id: str, paths: list[Path]) -> list[dict[str, Any]]:
    artifacts: list[dict[str, Any]] = []
    for path in paths:
        if not path.exists() or not path.is_file():
            continue
        artifacts.append({
            "name": path.name,
            "label": artifact_label(path),
            "size": path.stat().st_size,
            "url": f"/api/jobs/{job_id}/download/{path.name}",
        })
    return artifacts


def process_job(
    job_id: str,
    source_path: Path,
    original_name: str,
    mode: str,
    audio_format: str,
    model_name: str,
    language: str,
) -> None:
    try:
        update_job(job_id, status="processing", progress=5, message="Preparazione file")
        settings = read_settings()
        stamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        base_name = safe_stem(Path(original_name).stem)
        job_dir = output_root() / f"{stamp}_{base_name}_{job_id[:8]}"
        job_dir.mkdir(parents=True, exist_ok=False)
        update_job(job_id, output_dir=str(job_dir))

        update_job(job_id, progress=15, message="Estrazione audio")
        audio_path = job_dir / f"{base_name}.{audio_format}"
        run_command(audio_command(source_path, audio_path, audio_format), job_id)
        artifacts = [audio_path]

        if mode == "transcribe":
            if not whisper_available():
                raise RuntimeError("Whisper locale non è installato. Riesegui INSTALLA_MEDIA_LAB.bat.")
            update_job(job_id, progress=45, message="Preparazione audio per Whisper")
            whisper_input = job_dir / f"{base_name}_whisper.wav"
            run_command(whisper_input_command(audio_path, whisper_input), job_id)
            update_job(job_id, progress=60, message=f"Caricamento modello Whisper {model_name}")
            model = load_whisper_model(model_name)
            update_job(job_id, progress=72, message="Trascrizione locale in corso")
            transcribe_kwargs: dict[str, Any] = {"fp16": False, "verbose": False}
            if language != "auto":
                transcribe_kwargs["language"] = language
            result = model.transcribe(str(whisper_input), **transcribe_kwargs)
            update_job(job_id, progress=90, message="Creazione TXT, DOCX e SRT")
            artifacts.extend(write_transcription_files(result, job_dir, base_name))
            whisper_input.unlink(missing_ok=True)

        if bool(settings.get("keep_original", False)):
            original_copy = job_dir / safe_filename(original_name)
            shutil.copy2(source_path, original_copy)
            artifacts.append(original_copy)

        update_job(
            job_id,
            status="completed",
            progress=100,
            message="Elaborazione completata",
            artifacts=serialize_artifacts(job_id, artifacts),
            completed_at=time.time(),
        )
        add_log(job_id, "Elaborazione completata correttamente.")
    except Exception as exc:
        update_job(
            job_id,
            status="failed",
            progress=100,
            message="Elaborazione non riuscita",
            error=str(exc),
            completed_at=time.time(),
        )
        add_log(job_id, f"ERRORE: {exc}")
    finally:
        source_path.unlink(missing_ok=True)


def open_in_explorer(path: Path) -> None:
    path = path.resolve()
    if os.name == "nt":
        os.startfile(str(path))  # type: ignore[attr-defined]
    elif sys_platform() == "darwin":
        subprocess.Popen(["open", str(path)])
    else:
        subprocess.Popen(["xdg-open", str(path)])


def sys_platform() -> str:
    import sys
    return sys.platform


@app.get("/")
def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/api/health")
def health() -> dict[str, Any]:
    ffmpeg = ffmpeg_path()
    return {
        "status": "ok",
        "name": APP_NAME,
        "version": APP_VERSION,
        "ffmpeg": bool(ffmpeg),
        "ffmpeg_path": ffmpeg,
        "whisper": whisper_available(),
        "output_dir": str(output_root()),
        "queue_size": sum(1 for job in _jobs.values() if job.get("status") in {"queued", "processing"}),
    }


@app.get("/api/settings")
def get_settings() -> dict[str, Any]:
    settings = read_settings()
    settings["output_dir"] = str(output_root())
    return settings


@app.post("/api/settings")
def save_settings(payload: dict[str, Any]) -> dict[str, Any]:
    current = read_settings()
    raw_output = payload.get("output_dir", current["output_dir"])
    output_value = str(raw_output or current["output_dir"]).strip()
    output_path = Path(output_value).expanduser()
    try:
        output_path.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Cartella risultati non valida: {exc}") from exc

    audio_format = str(payload.get("default_audio_format", current.get("default_audio_format", "mp3"))).lower()
    model_name = str(payload.get("default_whisper_model", current.get("default_whisper_model", "base"))).lower()
    language = str(payload.get("default_language", current.get("default_language", "it"))).lower()
    if audio_format not in AUDIO_FORMATS:
        audio_format = "mp3"
    if model_name not in WHISPER_MODELS:
        model_name = "base"
    if language not in LANGUAGES:
        language = "it"

    current.update({
        "output_dir": str(output_path.resolve()),
        "keep_original": bool(payload.get("keep_original", current.get("keep_original", False))),
        "default_audio_format": audio_format,
        "default_whisper_model": model_name,
        "default_language": language,
    })
    write_settings(current)
    return current


@app.post("/api/jobs")
async def create_job(
    file: UploadFile = File(...),
    mode: str = Form("audio"),
    audio_format: str = Form("mp3"),
    model: str = Form("base"),
    language: str = Form("it"),
) -> dict[str, Any]:
    ensure_directories()
    mode = mode.lower().strip()
    audio_format = audio_format.lower().strip()
    model = model.lower().strip()
    language = language.lower().strip()

    if mode not in {"audio", "transcribe"}:
        raise HTTPException(status_code=400, detail="Modalità non valida")
    if audio_format not in AUDIO_FORMATS:
        raise HTTPException(status_code=400, detail="Formato audio non valido")
    if model not in WHISPER_MODELS:
        raise HTTPException(status_code=400, detail="Modello Whisper non valido")
    if language not in LANGUAGES:
        raise HTTPException(status_code=400, detail="Lingua non valida")

    original_name = safe_filename(file.filename or "media")
    extension = Path(original_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Formato file non supportato: {extension or 'senza estensione'}")

    job_id = uuid.uuid4().hex
    incoming_path = INCOMING_DIR / f"{job_id}_{original_name}"
    try:
        with incoming_path.open("wb") as destination:
            while chunk := await file.read(1024 * 1024):
                destination.write(chunk)
    finally:
        await file.close()

    if incoming_path.stat().st_size == 0:
        incoming_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="Il file caricato è vuoto")

    now = time.time()
    job = {
        "id": job_id,
        "filename": original_name,
        "mode": mode,
        "audio_format": audio_format,
        "model": model,
        "language": language,
        "status": "queued",
        "progress": 0,
        "message": "In coda",
        "error": None,
        "artifacts": [],
        "logs": ["File ricevuto dal computer locale."],
        "created_at": now,
        "updated_at": now,
        "output_dir": None,
    }
    with _jobs_lock:
        _jobs[job_id] = job

    _executor.submit(
        process_job,
        job_id,
        incoming_path,
        original_name,
        mode,
        audio_format,
        model,
        language,
    )
    return {"job_id": job_id, "status": "queued"}


@app.get("/api/jobs")
def list_jobs() -> dict[str, Any]:
    with _jobs_lock:
        jobs = [dict(job) for job in _jobs.values()]
    jobs.sort(key=lambda item: float(item.get("created_at", 0)), reverse=True)
    return {"jobs": jobs[:50]}


@app.get("/api/jobs/{job_id}")
def job_status(job_id: str) -> dict[str, Any]:
    return get_job(job_id)


@app.get("/api/jobs/{job_id}/download/{filename}")
def download_artifact(job_id: str, filename: str) -> FileResponse:
    job = get_job(job_id)
    output_dir_value = job.get("output_dir")
    if not output_dir_value:
        raise HTTPException(status_code=404, detail="Cartella risultati non disponibile")
    output_dir_path = Path(str(output_dir_value)).resolve()
    candidate = (output_dir_path / Path(filename).name).resolve()
    if candidate.parent != output_dir_path or not candidate.exists() or not candidate.is_file():
        raise HTTPException(status_code=404, detail="File risultato non trovato")
    return FileResponse(candidate, filename=candidate.name)


@app.post("/api/jobs/{job_id}/open-folder")
def open_job_folder(job_id: str) -> dict[str, Any]:
    job = get_job(job_id)
    output_dir_value = job.get("output_dir")
    if not output_dir_value:
        raise HTTPException(status_code=404, detail="Cartella risultati non disponibile")
    folder = Path(str(output_dir_value))
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Cartella risultati non trovata")
    open_in_explorer(folder)
    return {"opened": True, "path": str(folder)}


@app.post("/api/open-output")
def open_output_folder() -> dict[str, Any]:
    folder = output_root()
    open_in_explorer(folder)
    return {"opened": True, "path": str(folder)}


def open_browser() -> None:
    try:
        webbrowser.open(f"http://{HOST}:{PORT}")
    except Exception:
        pass


def main() -> None:
    ensure_directories()
    PID_FILE.write_text(str(os.getpid()), encoding="ascii")
    timer = threading.Timer(1.2, open_browser)
    timer.daemon = True
    timer.start()
    try:
        uvicorn.run(app, host=HOST, port=PORT, log_level="info")
    finally:
        PID_FILE.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
